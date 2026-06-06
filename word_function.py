from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt, RGBColor
from typing import Optional, Union
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_UNDERLINE, WD_COLOR_INDEX ,WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement , parse_xml
from docx.oxml.ns import qn , nsdecls
from docx.shared import Cm

""" ---------- تابع جادویی برای راست‌به‌چپ کردن کل سند ----------"""
def force_document_rtl(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        for b in sectPr.xpath('./w:bidi'):
            sectPr.remove(b)
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

    style = doc.styles['Normal']
    pPr = style._element.get_or_add_pPr()

    # حذف تنظیمات قبلی
    for jc in pPr.xpath('./w:jc'):
        pPr.remove(jc)
    for b in pPr.xpath('./w:bidi'):
        pPr.remove(b)

    # تراز پیش‌فرض واقعی
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


# ---------- تابع راست نویس پاراگراف ----------
def set_paragraph_rtl(p):
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    

#=============== set font,size -تنظیم Run برای متون RTL (فارسی/عربی) با قابلیت‌های سفارشی متعدد
# set_run_rtl(run,font_name="tahoma",size=18,bold=True,
#italic=True,underline=True,color_rgb="FF0000")
def set_run_rtl(
    run,
    font_name: str = "B Nazanin",
    size: Optional[Union[int, float]] = 13,
    bold: Optional[bool] = False,
    italic: Optional[bool] = False,
    underline: Optional[Union[bool, WD_UNDERLINE]] = False,
    color_rgb: Optional[str] = None,
    highlight_color: Optional[WD_COLOR_INDEX] = None,
    cs_font: Optional[str] = None,
 ):
        
    # تنظیمات پایه
    if font_name:
        run.font.name = font_name
    
    # تنظیم سایز فونت (به روش صحیح)
    if size is not None:
        # روش اول: تنظیم از طریق property معمولی
        run.font.size = Pt(size)
        
        # روش دوم: تنظیم مستقیم XML (برای اطمینان)
        r = run._element
        rPr = r.get_or_add_rPr()
        
        # حذف تگ‌های قدیمی سایز
        for sz_tag in ['w:sz', 'w:szCs']:
            old_sz = rPr.find(qn(sz_tag))
            if old_sz is not None:
                rPr.remove(old_sz)
        
        # اضافه کردن تگ جدید سایز (نصف مقدار Pt برای ذخیره در XML)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size * 2)))  # تبدیل Pt به Half-points
        rPr.append(sz)
        
        sz_cs = OxmlElement('w:szCs')
        sz_cs.set(qn('w:val'), str(int(size * 2)))
        rPr.append(sz_cs)
    
    if bold is not None:
        run.font.bold = bold
    
    if italic is not None:
        run.font.italic = italic
    
    if underline:
        run.font.underline = underline
    
    if color_rgb:
        red = int(color_rgb[0:2], 16)
        green = int(color_rgb[2:4], 16)
        blue = int(color_rgb[4:6], 16)
        run.font.color.rgb = RGBColor(red, green, blue)
    
    if highlight_color:
        run.font.highlight_color = highlight_color
    
    # عنصر XML زیرین
    r = run._element
    rPr = r.get_or_add_rPr()
    
    # تنظیم فونت برای متون پیچیده (فارسی)
    # حذف تگ rFonts قدیمی اگر وجود دارد
    old_rFonts = rPr.find(qn('w:rFonts'))
    if old_rFonts is not None:
        rPr.remove(old_rFonts)
    
    rFonts = OxmlElement('w:rFonts')
    if font_name:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    if cs_font:
        rFonts.set(qn('w:cs'), cs_font)
    elif font_name:
        rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    
    # فعال کردن RTL برای خود کلمات
    # حذف تگ rtl قدیمی اگر وجود دارد
    old_rtl = rPr.find(qn('w:rtl'))
    if old_rtl is not None:
        rPr.remove(old_rtl)
    
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)

def add_persian_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # حذف موارد قبلی
    for tag in ['jc', 'bidiVisual', 'tblInd', 'tblW', 'tblLayout']:
        for el in tblPr.xpath(f'./w:{tag}'):
            tblPr.remove(el)

    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    tblPr.append(jc)

    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)

    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9000')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # عرض ستون‌ها
    widths = [Cm(4), Cm(3), Cm(3)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            set_paragraph_rtl(p)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i == 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(str(val))
            set_run_rtl(run, size=12, bold=(i == 0))
    # غیرفعال کردن AutoFit (خیلی مهم!)
    table.allow_autofit = False
    # تعیین عرض ثابت برای جدول
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9000')   # معادل ~15 سانتی‌متر
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

     

    return table

def set_table_rtl(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)

def set_table_column_widths(table, widths_cm):
    """
    widths_cm: لیست عرض ستون‌ها بر حسب Cm
    مثال: [Cm(2), Cm(6), Cm(3), ...]
    """
    # 1) خاموش کردن AutoFit
    table.allow_autofit = False

    # 2) fixed layout
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # 3) تنظیم عرض ستون‌ها روی همه ردیف‌ها
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            if i >= len(row.cells):
                continue
            cell = row.cells[i]

            # روش سطح بالا
            cell.width = width

            # روش XML برای اطمینان
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)

            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(int(width.twips)))  # تبدیل Cm به twips

def set_cell_background(cell, color_hex ):
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex)
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_cell_text(cell, text, font_name="Tahoma", size=9, bold=False, align=WD_PARAGRAPH_ALIGNMENT.CENTER):
    # پاک کردن محتوای قبلی (اگر لازم است)
    for p in cell.paragraphs:
        p.clear()
        
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
      
    run = p.add_run(str(text))
    set_run_rtl(run, font_name=font_name, size=size, bold=bold)
    # فعال کردن RTL برای این پاراگراف
    set_paragraph_rtl(p)
    p.alignment = align
    return p

def set_cell_vertical_alignment(cell, alignment="center"):
    """
    Vertical align inside a table cell.
    alignment: 'top' | 'center' | 'bottom'
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # remove existing vAlign (avoid duplicates)
    for child in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(child)

    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), alignment)
    tcPr.append(vAlign)
    
# تابع برای تبدیل اعداد انگلیسی به فارسی و سه رقم جدا کردن
def num_rial(number):

    if number is None or number == "":
        return "۰"
    
    try:
        # ۱. تبدیل ورودی به عدد صحیح (در صورتی که به صورت رشته یا اعشار باشد)
        num_val = int(float(number))
        
        # ۲. سه رقم جدا کردن (خروجی این یک رشته انگلیسی است)
        formatted_str = "{:,}".format(num_val)
        
        # ۳. تعریف جدول ترجمه برای تبدیل اعداد به فارسی
        translation_table = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
        
        # ۴. اعمال ترجمه روی رشته (دقت کنید: formatted_str.translate نه num_rial.translate)
        return formatted_str.translate(translation_table)
        
    except (ValueError, TypeError):
        # در صورت بروز خطا در تبدیل، خود مقدار را برگردان
        return str(number)

def float_rial(number, decimal_places=2):
    """
    اعداد اعشاری را با سه رقم جدا کردن و اعشار مشخص، به فارسی برمی‌گرداند.
    مثال: 1250000.567 -> ۱,۲۵۰,۰۰۰.۵۷
    """
    if number is None:
        return "۰"
    
    try:
        # تبدیل به عدد اعشاری برای اطمینان
        num_val = float(number)
        
        # قالب‌بندی با تعداد اعشار مشخص
        # format_str خروجی به صورت "1,250,000.57" می‌دهد
        format_str = "{:,.{}f}".format(num_val, decimal_places)
        
        # جدول تبدیل اعداد انگلیسی به فارسی
        # نکته: ممیز انگلیسی (.) را به ممیز فارسی (٫) تبدیل می‌کنیم
        trans_table = str.maketrans("0123456789.", "۰۱۲۳۴۵۶۷۸۹٫")
        
        return format_str.translate(trans_table)
        
    except (ValueError, TypeError):
        return str(number)


# تابع تبدیل عدد به حروف (نوشتار فارسی)
def num_word(num):
    if num == 0:
        return "صفر"

    ones = ["", "یک", "دو", "سه", "چهار", "پنج", "شش", "هفت", "هشت", "نه"]
    tens = ["", "ده", "بیست", "سی", "چهل", "پنجاه", "شصت", "هفتاد", "هشتاد", "نود"]
    teens = ["ده", "یازده", "دوازده", "سیزده", "چهارده", "پانزده", "شانزده", "هفده", "هجده", "نوزده"]
    hundreds = ["", "صد", "دویست", "سیصد", "چهارصد", "پانصد", "ششصد", "هفتصد", "هشتصد", "نهصد"]
    thousands = ["", "هزار", "میلیون", "میلیارد", "تریلیون"]

    def convert_section(n):
        res = ""
        h = n // 100
        t = (n % 100) // 10
        o = n % 10

        if h > 0:
            res += hundreds[h]
        
        if t == 1:
            if res: res += " و "
            res += teens[o]
        else:
            if t > 0:
                if res: res += " و "
                res += tens[t]
            if o > 0:
                if res: res += " و "
                res += ones[o]
        return res

    words = []
    unit_idx = 0
    
    temp_num = abs(num)
    while temp_num > 0:
        section = temp_num % 1000
        if section > 0:
            section_word = convert_section(section)
            if unit_idx > 0:
                section_word += " " + thousands[unit_idx]
            words.insert(0, section_word)
        temp_num //= 1000
        unit_idx += 1

    result = " و ".join(words)
    return "منفی " + result if num < 0 else result

"""
# --- مثال استفاده ---
input_num = 1250750

# نمایش خروجی
print(f"عدد اصلی: {input_num}")
print(f"نمایش عددی فارسی: {format_and_farsify(input_num)}")
print(f"نمایش حروفی فارسی: {number_to_persian_words(input_num)}")

"""
""" 
# رنگ‌های پایه
color_rgb="0000FF"   # آبی (Blue)
color_rgb="FF0000"   # قرمز (Red)
color_rgb="00FF00"   # سبز (Green)
color_rgb="000000"   # مشکی (Black)
color_rgb="FFFFFF"   # سفید (White)

# رنگ‌های دیگر
color_rgb="FFFF00"   # زرد (Yellow)
color_rgb="FF00FF"   # صورتی/مغزى (Magenta)
color_rgb="00FFFF"   # فیروزه‌ای (Cyan)
color_rgb="800080"   # بنفش (Purple)
color_rgb="FFA500"   # نارنجی (Orange)
color_rgb="808080"   # خاکستری (Gray)
color_rgb="000080"   # آبی سیر (Navy)
color_rgb="ADD8E6"   # آبی روشن (Light Blue)
color_rgb="1E90FF"   # آبی آسمانی (Dodger Blue)
color_rgb="4169E1"   # آبی سلطنتی (Royal Blue)
"""

# ---------- اجرای برنامه ----------

#doc = Document()

# اعمال RTL به کل ساختار فایل
#force_document_rtl(doc)

# افزودن تیتر
"""p_title = doc.add_paragraph()
p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_title = p_title.add_run("گزارش کاملاً راست‌به‌چپ")
set_run_rtl(run_title, size=20, bold=True)

# افزودن متن با تراز Justify
p_text = doc.add_paragraph()
p_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run_text = p_text.add_run(
    "در این نسخه، کل بدنه فایل ورد (Section Properties) تغییر یافته است. "
    "حتی اگر شما در فایل نهایی کلیک کنید و اینتر بزنید، مکان‌نما (Cursor) "
    "به صورت خودکار در سمت راست قرار می‌گیرد چون کل سند RTL شده است."
)

set_run_rtl(run_text,font_name="tahoma",size=10,bold=True,
italic=False,underline=False,color_rgb="000080")


# افزودن جدول
data = [
    ["محصول", "کد", "موجودی"],
    ["لپ‌تاپ", "A10", "۵ عدد"],
    ["مانیتور", "B20", "۱۲ عدد"]
]
# افزودن خط خالی برای جداسازی پاراگراف قبلی و جدول
p_before_table = doc.add_paragraph()
add_persian_table(doc, data)

# ذخیره سازی ایمن
try:
    doc.save("rtl_report.docx")
    print("فایل با موفقیت و با ساختار RTL کامل ساخته شد.")
except PermissionError:
    print("خطا: فایل در ورد باز است. آن را ببندید.")
"""