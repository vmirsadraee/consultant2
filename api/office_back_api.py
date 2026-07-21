from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm , Pt 
from docx.enum.table import WD_ROW_HEIGHT_RULE # برای کنترل دقیق ارتفاع
from .word_function import ( set_table_rtl , add_cell_text , add_persian_table , set_run_rtl
     , set_paragraph_rtl , force_document_rtl , set_table_column_widths , set_cell_background
     , num_rial , float_rial , num_word , set_cell_vertical_alignment)

# ---------- RTL helper ----------

import json

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def parse_table(value):
    if isinstance(value, str):
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            return None
    if isinstance(value, dict):
        return value
    return None


# -------------------------
# ساخت جدول ورد
# -------------------------
def render_table(doc, table_data):
    if not table_data:
        return

    title = table_data.get("title", "")
    columns = table_data.get("columns", [])
    rows = table_data.get("rows", [])
    total = table_data.get("total", None)
    currency = table_data.get("currency", "")
    #h = doc.add_heading(title, level=2)
    #set_paragraph_rtl(h)

    h_paragraph=doc.add_paragraph() 
    h_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    h =h_paragraph.add_run(title)
    set_run_rtl(h,font_name="tahoma",size=10,bold=True,
               italic=False,underline=False,color_rgb="000080")
    # ==========================
    # محاسبه تعداد ستون واقعی مورد نیاز بر اساس داده‌های موجود
    max_cols = len(columns)
    for r in rows:
        kind = r.get("kind")
        if kind == "data1":
            max_cols = max(max_cols, 6)
        elif kind == "data2":
            max_cols = max(max_cols, 7)
        elif kind == "data3":
            max_cols = max(max_cols, 8)
            
    table = doc.add_table(rows=1, cols=max_cols)
    table.style = "Table Grid"
    table.autofit = False
    set_table_rtl(table)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # مثال: عرض ستون‌ها بر اساس نوع جدول شما
    # برای data1 (6 ستون): سریال، شرح، مقدار، فی، ضریب، جمع
    if max_cols == 6:
        widths = [Cm(2), Cm(6), Cm(2.5), Cm(2.5), Cm(1.5), Cm(2.5)]
    elif max_cols == 7:
        widths = [Cm(2), Cm(6), Cm(2.5), Cm(2.0), Cm(1.5), Cm(1.5), Cm(2.5)]
    elif max_cols == 8:
        widths = [Cm(2), Cm(6.0), Cm(2.5), Cm(1.8),Cm(1.5), Cm(1.5), Cm(1.5), Cm(2.5)]
    else:
        # حالت ساده: همه مساوی
        widths = [Cm(2.5)] * max_cols
        
    set_table_column_widths(table, widths)
    # ==========================
     # تنظیم هدرها
    table.rows[0].height = Pt(25)
    header_cells = table.rows[0].cells

    # درست:
    for cell in header_cells:set_cell_vertical_alignment(cell, "center")

    for i, col_text in enumerate(columns):
        # گرفتن یا ایجاد پاراگراف در سلول
        cell = header_cells[i]             
        set_cell_background(cell, "4169E1" )  # زرد
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        
        run = p.add_run(str(col_text))     # ایجاد Run جدید در پاراگراف
               
        set_run_rtl(run, font_name="Tahoma", size=11, bold=True,     # اعمال استایل
                    italic=False, underline=False, color_rgb="000000")     
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for r in rows:

        kind = r.get("kind")

        # ----------------
        # data row
        # ----------------
        if kind == "data1":  # before and after
            #cells = table.add_row().cells
            row_obj = table.add_row() # سطر را در یک متغیر ذخیره کنید
            cells =  row_obj.cells
            # بررسی اینکه آیا این سطر یک عنوان است یا داده عادی
            if r.get("title_1") in [0, ""]:
                row_obj.height = Pt(25)
                add_cell_text(cells[0], r.get("serial") or "", bold=True ,align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                merged_cell = cells[1].merge(cells[5])
                add_cell_text(merged_cell, r.get("subject") or "", size=11, bold=True , 
                         align=WD_PARAGRAPH_ALIGNMENT.LEFT)
               
                set_cell_background(merged_cell, "00FFFF" )  # زرد  
                set_cell_background(cells[0], "00FFFF")  # زرد 
           
           
           
            else:
                 #def add_cell_text(cell, text, font_name="Tahoma",
                 #size=9, bold=False, align=WD_PARAGRAPH_ALIGNMENT.CENTER):
                add_cell_text(cells[0], r.get("serial") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                add_cell_text(cells[1], r.get("subject") or "", align=WD_PARAGRAPH_ALIGNMENT.LEFT)
                add_cell_text(cells[2], r.get("deliver") or "", size=7 , align=WD_PARAGRAPH_ALIGNMENT.CENTER)
               
                unit_value=r.get("unitprice") or ""
                add_cell_text(cells[3], num_rial(unit_value) , align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                
                a=r.get("coefficienta") or ""
                if not r.get("totalprice")== 0 or '' : 
                    add_cell_text(cells[4], float_rial(a), align=WD_PARAGRAPH_ALIGNMENT.CENTER)
               
                total_value=r.get("totalprice") or ""
                add_cell_text(cells[5], num_rial(total_value) or "", size=10 , align=WD_PARAGRAPH_ALIGNMENT.CENTER)

        elif kind == "data2":  # durring
            cells = table.add_row().cells
            # بررسی اینکه آیا این سطر یک عنوان است یا داده عادی
            if r.get("title_1") in [0, ""]:
                add_cell_text(cells[0], r.get("serial") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                merged_cell = cells[1].merge(cells[6])
                add_cell_text(merged_cell, r.get("subject") or "", align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            else:
                 # 2. اگر مرج نیست، تک‌تک بنویس
                add_cell_text(cells[0], r.get("serial") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                add_cell_text(cells[1], r.get("subject") or "", align=WD_PARAGRAPH_ALIGNMENT.LEFT)
                add_cell_text(cells[2], r.get("deliver") or "", size=7, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                add_cell_text(cells[3], r.get("unitprice") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                if not r.get("totalprice")== 0 or '' : 
                    add_cell_text(cells[4], r.get("coefficienta") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                if not r.get("totalprice")== 0 or '' : 
                    add_cell_text(cells[5], r.get("coefficientb") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                add_cell_text(cells[6], r.get("totalprice") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)

            
        elif kind == "data3":  # case
            
            cells = table.add_row().cells
            # بررسی اینکه آیا این سطر یک عنوان است یا داده عادی
            if r.get("title_1") in [0, ""]:
                add_cell_text(cells[0], r.get("serial") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                merged_cell = cells[1].merge(cells[7])
                add_cell_text(merged_cell, r.get("subject") or "", align=WD_PARAGRAPH_ALIGNMENT.LEFT)
            else:
                 # 2. اگر مرج نیست، تک‌تک بنویس
                add_cell_text(cells[0], r.get("serial") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                add_cell_text(cells[1], r.get("subject") or "", align=WD_PARAGRAPH_ALIGNMENT.LEFT)
                add_cell_text(cells[2], r.get("deliver") or "" , size=7,  align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                add_cell_text(cells[3], r.get("unitprice") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                if not r.get("totalprice")== 0 or '' :                   
                       add_cell_text(cells[4], r.get("coefficienta") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                
                if not r.get("totalprice")== 0 or '' : 
                       add_cell_text(cells[5], r.get("coefficientb") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
               
                if not r.get("totalprice")== 0 or '' : 
                       add_cell_text(cells[6], r.get("coefficientc") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
                add_cell_text(cells[7], r.get("totalprice") or "", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    set_table_column_widths(table, widths)  

    if total is not None:
        p = doc.add_paragraph(f"جمع کل: {num_rial(total)} {currency}")
        set_paragraph_rtl(p)

        p = doc.add_paragraph()
        p.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run(f"جمع کل: {num_word(total)} {currency}")     # ایجاد Run جدید در پاراگراف      
        set_run_rtl(run, font_name="Tahoma", size=12, bold=True,     # اعمال استایل
                    italic=False, underline=False, color_rgb="000080")
        doc.add_page_break()


# -------------------------
# endpoint
# -------------------------
@app.post("/api/report/generate")
async def generate_report(payload: dict):

    tables = payload.get("tables", {})

    before = parse_table(tables.get("before"))
    durring = parse_table(tables.get("durring"))
    case = parse_table(tables.get("case"))
    # site = parse_table(tables.get("site"))
    after = parse_table(tables.get("after"))
    # support = parse_table(tables.get("support"))

    doc = Document()
    force_document_rtl(doc)   
    h_paragraph=doc.add_paragraph() 
    h_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    h =h_paragraph.add_run("گزارش پایش")
    set_run_rtl(h,font_name="tahoma",size=24,bold=True,
               italic=False,underline=False,color_rgb="FF0000")


    render_table(doc, before)
    render_table(doc, durring)
    render_table(doc, case)
    # render_table(doc, site)
    render_table(doc, after)
    # render_table(doc, support)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="report.docx"'},
    )

print("========== ROUTES ==========")
for r in app.routes:
    print(r.path, r.methods)
print("============================")
